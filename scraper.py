import os
import time
import csv
from playwright.sync_api import sync_playwright
from docx import Document

def create_status_doc(results, summary):
    doc = Document()
    doc.add_heading('Kolkata Lead Hunter: Master Market Audit', 0)
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(f"Total Listings Scanned: {summary['total']}")
    doc.add_paragraph(f"High-Value Targets Added: {summary['added']}")
    doc.add_paragraph(f"Skipped (Known/Has Site): {summary['skipped']}")
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Company Name'
    hdr_cells[1].text = 'Status'
    hdr_cells[2].text = 'Reason'
    for item in results:
        row_cells = table.add_row().cells
        row_cells[0].text = item['name']
        row_cells[1].text = item['status']
        row_cells[2].text = item['reason']
    doc.save('status_report.docx')

def hunt_leads():
    query = os.getenv("SEARCH_QUERY", "HR Consultancy in Kolkata")
    search_url = f"https://www.google.com/maps/search/{query.replace(' ', '+')}/@22.5726,88.3639,13z"
    
    history_file = 'history.txt'
    if not os.path.exists(history_file):
        open(history_file, 'w').close()
    with open(history_file, 'r') as f:
        seen_names = set(line.strip() for line in f)

    processing_log = []
    stats = {"total": 0, "added": 0, "skipped": 0}
    leads_for_csv = []
    new_seen = []

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        context = browser.new_context(viewport={'width': 1920, 'height': 1080})
        page = context.new_page()
        
        page.goto(search_url)
        page.wait_for_selector('div[role="article"]', timeout=15000)

        # Deep Scroll Logic
        last_count = 0
        for _ in range(25):
            panel = page.locator('div[role="feed"]')
            if panel.count() > 0:
                panel.hover()
                page.mouse.wheel(0, 10000)
                time.sleep(4)
            current_count = page.locator('div[role="article"]').count()
            if current_count == last_count or current_count > 150: break
            last_count = current_count

        listings = page.locator('div[role="article"]').all()
        
        for listing in listings:
            stats["total"] += 1
            name = listing.get_attribute("aria-label") or "Unknown"
            
            if name in seen_names:
                processing_log.append({"name": name, "status": "SKIPPED", "reason": "Already in history."})
                stats["skipped"] += 1
                continue

            if listing.locator('a[data-item-id="authority"]').count() > 0:
                processing_log.append({"name": name, "status": "SKIPPED", "reason": "Already has a website."})
                stats["skipped"] += 1
                new_seen.append(name)
                continue

            # Deep Data Extraction
            listing.click()
            time.sleep(2.5) # Extra time for sidebar data to load

            if page.get_by_text("Add website").is_visible():
                # --- NEW MASTER DATA POINTS ---
                
                # 1. Phone
                phone = "N/A"
                phone_el = page.locator('button[data-item-id^="phone:tel"]').first
                if phone_el.is_visible():
                    phone = phone_el.get_attribute("data-item-id").replace("phone:tel:", "")

                # 2. Address
                address = "N/A"
                addr_el = page.locator('button[data-item-id="address"]').first
                if addr_el.is_visible():
                    address = addr_el.get_attribute("aria-label").replace("Address: ", "")

                # 3. Category
                category = "Business"
                cat_el = page.locator('button[data-item-id="category"]').first
                if cat_el.is_visible():
                    category = cat_el.inner_text()

                # 4. Rating & Reviews
                rating = "N/A"
                reviews = "0"
                try:
                    rating_el = page.locator('span.ceNzR').first # Google Maps Rating Class
                    if rating_el.is_visible():
                        rating = rating_el.inner_text()
                except: pass

                # 5. Maps Link
                maps_url = page.url

                processing_log.append({"name": name, "status": "ADDED", "reason": "Confirmed Target Lead."})
                leads_for_csv.append({
                    "Name": name, 
                    "Phone": phone, 
                    "Address": address, 
                    "Category": category, 
                    "Rating": rating,
                    "MapsLink": maps_url
                })
                stats["added"] += 1
            else:
                processing_log.append({"name": name, "status": "SKIPPED", "reason": "No 'Add website' option."})
                stats["skipped"] += 1
            
            new_seen.append(name)
            page.keyboard.press("Escape")

        # Save Logic
        with open(history_file, 'a') as f:
            for n in new_seen: f.write(n + '\n')
        
        with open('leads.csv', 'w', newline='', encoding='utf-8') as f:
            fieldnames = ["Name", "Phone", "Address", "Category", "Rating", "MapsLink"]
            writer = csv.DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(leads_for_csv)

        create_status_doc(processing_log, stats)
        browser.close()

if __name__ == "__main__":
    hunt_leads()
